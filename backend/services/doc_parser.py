"""
Document parser — extracts headings, text, and tables from .docx files,
and page-by-page text from .pdf files. Both produce ``DocSection`` lists
so the shared ``doc_chunker`` can process either source unchanged.
"""
import logging
import os
import re
from dataclasses import dataclass, field

logger = logging.getLogger("projecthub.doc_parser")

# Per-page text cap for PDFs — keeps a single huge page from blowing the
# LLM context once chunked. ~2K tokens of prose.
_PDF_PER_PAGE_CHAR_CAP = 8000


@dataclass
class DocSection:
    heading: str
    heading_level: int
    heading_path: str  # e.g. "3. Architektur > 3.2 Backend > 3.2.1 Schema"
    text: str
    tables: list[str] = field(default_factory=list)
    has_images: bool = False
    char_count: int = 0

    def __post_init__(self):
        self.char_count = len(self.text) + sum(len(t) for t in self.tables)


def parse_docx(file_path: str) -> list[DocSection]:
    """Parse a .docx file into sections based on headings."""
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error("Failed to open docx %s: %s", file_path, e)
        return []

    sections: list[DocSection] = []
    heading_stack: list[str] = []  # Current heading path
    current_heading = "Einleitung"
    current_level = 0
    current_text_parts: list[str] = []
    current_tables: list[str] = []
    current_has_images = False

    def _flush_section():
        """Save current accumulated text as a section."""
        nonlocal current_text_parts, current_tables, current_has_images
        text = "\n".join(current_text_parts).strip()
        if text or current_tables:
            path = " > ".join(heading_stack) if heading_stack else current_heading
            sections.append(DocSection(
                heading=current_heading,
                heading_level=current_level,
                heading_path=path,
                text=text,
                tables=current_tables[:],
                has_images=current_has_images,
            ))
        current_text_parts = []
        current_tables = []
        current_has_images = False

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        # Detect headings
        if style_name.startswith("Heading"):
            _flush_section()
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 1

            heading_text = para.text.strip()
            if not heading_text:
                continue

            # Update heading stack
            while len(heading_stack) >= level:
                heading_stack.pop()
            heading_stack.append(heading_text)

            current_heading = heading_text
            current_level = level

        else:
            # Regular paragraph
            text = para.text.strip()
            if text:
                current_text_parts.append(text)

            # Check for inline images
            for run in para.runs:
                if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    current_has_images = True
                if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
                    current_has_images = True

    # Parse tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if not rows:
            continue

        # Convert to markdown table
        md_lines = []
        # Header
        md_lines.append("| " + " | ".join(rows[0]) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
        # Body
        for row in rows[1:]:
            # Pad row if shorter
            while len(row) < len(rows[0]):
                row.append("")
            md_lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")

        # Attach table to the last section or create a standalone one
        table_md = "\n".join(md_lines)
        if sections:
            sections[-1].tables.append(table_md)
            sections[-1].char_count += len(table_md)
        else:
            current_tables.append(table_md)

    # Flush last section
    _flush_section()

    logger.info("Parsed %s: %d sections", file_path, len(sections))
    return sections


def parse_pdf(file_path: str) -> list[DocSection]:
    """Parse a .pdf file into one ``DocSection`` per (non-empty) page.

    PDFs have no reliable heading structure, so each page becomes its own
    section ("Seite N"). The downstream ``chunk_sections`` then merges
    small pages and splits large ones — identical handling to .docx.

    Scanned PDFs (no extractable text layer) yield zero sections; the
    caller treats that as a parse with no content rather than an error.
    """
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError
    except ImportError as e:
        logger.error("pypdf not installed — cannot parse PDF %s: %s", file_path, e)
        return []

    try:
        reader = PdfReader(file_path)
    except PdfReadError as e:
        logger.error("Corrupt/unsupported PDF %s: %s", file_path, e)
        return []
    except Exception as e:
        logger.error("Failed to open PDF %s: %s", file_path, e)
        return []

    if reader.is_encrypted:
        # pypdf can often decrypt with an empty password.
        try:
            if not reader.decrypt(""):
                logger.warning("PDF %s is encrypted — skipping", file_path)
                return []
        except Exception:
            logger.warning("PDF %s is encrypted — skipping", file_path)
            return []

    doc_name = os.path.basename(file_path)
    sections: list[DocSection] = []
    for idx, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as e:
            logger.warning("Page %d extract failed in %s: %s", idx, doc_name, e)
            continue
        if not text:
            continue  # empty / scanned page — no signal
        if len(text) > _PDF_PER_PAGE_CHAR_CAP:
            text = text[:_PDF_PER_PAGE_CHAR_CAP]
        heading = f"Seite {idx}"
        sections.append(DocSection(
            heading=heading,
            heading_level=1,
            heading_path=f"{doc_name} > {heading}",
            text=text,
        ))

    logger.info("Parsed PDF %s: %d sections from %d pages",
                file_path, len(sections), len(reader.pages))
    return sections
